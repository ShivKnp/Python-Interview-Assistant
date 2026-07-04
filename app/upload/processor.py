"""
Document Processor — parses uploaded files (PDF, DOCX, TXT, MD, PY)
into LangChain Documents ready for embedding and indexing.
"""

from __future__ import annotations

import io
import re
from pathlib import Path

from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter

from app.config import get_settings
from app.observability.logger import logger

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md", ".py", ".rst"}


def _extract_pdf(content: bytes, filename: str) -> str:
    """Extract text from PDF bytes using pypdf."""
    try:
        import pypdf
        reader = pypdf.PdfReader(io.BytesIO(content))
        return "\n\n".join(
            page.extract_text() or "" for page in reader.pages
        ).strip()
    except ImportError:
        logger.warning("pypdf not installed — PDF extraction unavailable")
        return ""
    except Exception as exc:
        logger.error(f"PDF extraction failed for {filename}: {exc}")
        return ""


def _extract_docx(content: bytes, filename: str) -> str:
    """Extract text from DOCX bytes using python-docx."""
    try:
        import docx
        doc = docx.Document(io.BytesIO(content))
        return "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
    except ImportError:
        logger.warning("python-docx not installed — DOCX extraction unavailable")
        return ""
    except Exception as exc:
        logger.error(f"DOCX extraction failed for {filename}: {exc}")
        return ""


def _clean_text(text: str) -> str:
    """Normalize whitespace and remove null bytes."""
    text = text.replace("\x00", "")
    text = re.sub(r"\r\n", "\n", text)
    text = re.sub(r"\n{4,}", "\n\n\n", text)
    return text.strip()


def process_file(
    content: bytes,
    filename: str,
    user_id: str,
    doc_id: str,
) -> list[Document]:
    """
    Parse uploaded file bytes into chunked LangChain Documents.

    Args:
        content:  Raw file bytes.
        filename: Original filename (used for extension detection).
        user_id:  Owner's user identifier.
        doc_id:   UUID assigned to this upload (for deletion tracking).

    Returns:
        List of Document objects ready for embedding.
    """
    settings = get_settings()
    suffix = Path(filename).suffix.lower()

    if suffix not in SUPPORTED_EXTENSIONS:
        raise ValueError(
            f"Unsupported file type '{suffix}'. "
            f"Supported: {', '.join(SUPPORTED_EXTENSIONS)}"
        )

    # ── Extract raw text ──────────────────────────────────────────────────
    if suffix == ".pdf":
        raw_text = _extract_pdf(content, filename)
    elif suffix == ".docx":
        raw_text = _extract_docx(content, filename)
    else:
        # TXT, MD, PY, RST — decode as UTF-8
        try:
            raw_text = content.decode("utf-8", errors="replace")
        except Exception:
            raw_text = content.decode("latin-1", errors="replace")

    raw_text = _clean_text(raw_text)

    if not raw_text:
        raise ValueError(f"No extractable text found in '{filename}'")

    # ── Chunk ─────────────────────────────────────────────────────────────
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=settings.CHUNK_SIZE,
        chunk_overlap=settings.CHUNK_OVERLAP,
        separators=["\n\n\n", "\n\n", "\n", ". ", " "],
    )

    base_doc = Document(
        page_content=raw_text,
        metadata={
            "source": "user_doc",
            "filename": filename,
            "user_id": user_id,
            "doc_id": doc_id,
            "file_type": suffix.lstrip("."),
            "question_id": 0,
            "q_score": 0,
        },
    )
    chunks = splitter.split_documents([base_doc])

    logger.info_data(
        "Document processed",
        filename=filename,
        user_id=user_id,
        doc_id=doc_id,
        raw_chars=len(raw_text),
        chunks=len(chunks),
    )

    return chunks
